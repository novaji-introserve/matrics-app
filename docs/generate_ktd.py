"""Generate the iComply / MATRICS Knowledge Transfer Document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT = "Calibri"
BASE_PT = 12


# ─── helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **edges):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in edges.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para(doc, text, bold=False, size=BASE_PT, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 18, 2: 15, 3: 13}
    colors = {1: "1F3864", 2: "2E5F8A", 3: "365F91"}
    p = para(doc, text, bold=True, size=sizes.get(level, BASE_PT),
             color=colors.get(level, "000000"), space_before=12, space_after=4)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    indent = Inches(0.25 * level)
    p.paragraph_format.left_indent = indent + Inches(0.25)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(BASE_PT)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(BASE_PT)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(0)
    return p


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def simple_table(doc, headers, rows, header_bg="1F3864", header_fg="FFFFFF"):
    col_count = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=col_count)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT
        run.font.size = Pt(BASE_PT - 1)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(BASE_PT - 1)
    return t


# ─── architecture diagram helpers ─────────────────────────────────────────────

def arch_diagram_infra(doc):
    """Infrastructure diagram as a styled table."""
    para(doc, "Figure 1 – Infrastructure Overview", bold=True, size=10,
         color="555555", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)

    rows_def = [
        # layer label, components
        ("Internet / Client",  "Browser  ·  Mobile App  ·  REST Client"),
        ("TLS Termination",    "Nginx (HTTPS :443)  —  wildcard cert *.novajii.com"),
        ("Application Layer",  "Odoo 16  (4 workers, port 8069)     FastAPI (port 8001)"),
        ("Async / Events",     "Odoo Longpolling :8072   ·   Custom WS Server :8073   ·   Redis (queue_job, sessions, ORM cache)"),
        ("Data Layer",         "PostgreSQL 15  ←  PgBouncer (connection pool :5433)"),
        ("ETL / Integration",  "Apache SeaTunnel (REST :9080, cluster :5801)"),
        ("Storage",            "${STORAGE_PATH}/odoo16/  ·  /postgresql/  ·  /redis/"),
    ]

    layer_colors = ["2E5F8A", "365F91", "1F6E43", "1F6E43", "6B2D8A", "8A4B1F", "555555"]

    t = doc.add_table(rows=len(rows_def), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Inches(1.6), Inches(4.8)]

    for i, ((label, content), color) in enumerate(zip(rows_def, layer_colors)):
        row = t.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        _set_cell_bg(row.cells[0], color)
        r0 = row.cells[0].paragraphs[0].add_run(label)
        r0.font.name = FONT
        r0.font.size = Pt(10)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        bg = "F0F4FF" if i % 2 == 0 else "FAFAFA"
        _set_cell_bg(row.cells[1], bg)
        r1 = row.cells[1].paragraphs[0].add_run(content)
        r1.font.name = FONT
        r1.font.size = Pt(10)

    doc.add_paragraph()


def arch_diagram_modules(doc):
    """Module dependency diagram as a styled table."""
    para(doc, "Figure 2 – Custom Module Dependency Graph", bold=True, size=10,
         color="555555", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)

    layers = [
        ("Odoo 16 Base + OCA (base, mail, hr, bus, queue_job)", "2C3E50", 1),
        ("compliance_management  ← central hub (all modules depend on this)", "1A5276", 1),
        ("alert_management    |    regulatory_reports    |    etl_manager", "1F618D", 3),
        ("case_management  (depends on: alert_management)", "117A65", 1),
        ("nfiu_reporting  (depends on: case_management + regulatory_reports)", "6C3483", 1),
        ("internal_control  (depends on: alert_management + transaction_screening)", "7D6608", 1),
        ("access_control  ·  rule_book  ·  session_control  ·  (utility modules)", "616A6B", 1),
    ]

    t = doc.add_table(rows=len(layers), cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (text, color, _) in enumerate(layers):
        cell = t.rows[i].cells[0]
        _set_cell_bg(cell, color)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(text)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()


def arch_diagram_request_flow(doc):
    """HTTP request flow diagram."""
    para(doc, "Figure 3 – Request Flow", bold=True, size=10,
         color="555555", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)

    steps = [
        ("1", "Client", "HTTPS request to dev1.novajii.com", "2980B9"),
        ("2", "Nginx", "TLS termination → proxy_pass to upstream", "1A5276"),
        ("3a", "Odoo Workers", "/  →  Odoo HTTP (8069)", "1F6E43"),
        ("3b", "Odoo Longpolling", "/websocket, /longpolling/  →  port 8072", "1F6E43"),
        ("3c", "Custom WS", "/csv_import/ws  →  port 8073", "1F6E43"),
        ("3d", "FastAPI", "/api/*  →  FastAPI (8001)", "8E44AD"),
        ("4", "Database", "Odoo ORM → PgBouncer (5433) → PostgreSQL (5432)", "7D6608"),
        ("5", "Cache / Queue", "Redis: sessions, ORM cache, queue_job workers", "C0392B"),
    ]

    t = doc.add_table(rows=len(steps), cols=4)
    t.style = "Table Grid"
    for i, (step, component, detail, color) in enumerate(steps):
        row = t.rows[i]
        bg = "F8F9FA" if i % 2 == 0 else "FFFFFF"

        for j, (val, w) in enumerate([(step, Inches(0.35)), (component, Inches(1.3)),
                                       (detail, Inches(3.8)), ("", Inches(0.3))]):
            cell = row.cells[j]
            cell.width = w
            if j == 0:
                _set_cell_bg(cell, color)
                r = cell.paragraphs[0].add_run(val)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                _set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(val)
            r.font.name = FONT
            r.font.size = Pt(10)
            r.font.bold = (j == 1)
    doc.add_paragraph()


# ─── main document ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BASE_PT)

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────────────
    para(doc, "", space_before=20, space_after=0)
    para(doc, "iComply / MATRICS", bold=True, size=28, color="1F3864",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    para(doc, "Technical Knowledge Transfer Document", bold=True, size=16,
         color="2E5F8A", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "For New Developer Onboarding", size=13, color="555555",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Prepared by: Novaji Introserve Limited  |  April 2026",
         size=11, color="888888", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Confidential — Internal Use Only", size=10, color="CC0000",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    # ── 1. Project Overview ────────────────────────────────────────────────────
    heading(doc, "1. Project Overview")
    para(doc, (
        "iComply is a financial compliance management platform built on Odoo 16.0, "
        "designed for Nigerian financial institutions to meet AML/CFT (Anti-Money "
        "Laundering / Countering the Financing of Terrorism) regulatory requirements. "
        "The platform is deployed under the MATRICS brand and covers the full compliance "
        "lifecycle: customer KYC/KYB onboarding, risk assessment, transaction monitoring, "
        "alert management, case investigation, regulatory reporting (NFIU/STR/CTR), and "
        "internal control."
    ), space_after=8)

    heading(doc, "1.1  Business Context", level=2)
    bullet(doc, "Client: Nigerian commercial bank / financial institution")
    bullet(doc, "Regulatory frameworks: NFIU, CBN AML/CFT guidelines, FATF recommendations")
    bullet(doc, "Scale target: millions of customer records and transactions")
    bullet(doc, "Deployment environment: Ubuntu 22.04 LTS on cloud VM, Docker Compose stack")
    doc.add_paragraph()

    heading(doc, "1.2  Key Capabilities", level=2)
    rows = [
        ("Customer Management", "KYC/KYB onboarding, risk scoring, PEP/sanctions screening, EDD"),
        ("Transaction Monitoring", "Rule-based screening, alert generation, case workflows"),
        ("Risk Assessment", "Weighted scoring engine, risk universe, jurisdiction plans"),
        ("Regulatory Reporting", "NFIU STR/CTR, automated report generation and submission"),
        ("Case Management", "Investigation workflows, exception handling, audit trail"),
        ("Dashboard & Analytics", "Real-time charts, materialized views, KPI cards (OWL)"),
        ("ETL / Data Integration", "Apache SeaTunnel pipelines, FastAPI fraud-detection service"),
        ("Security", "Multi-layer SQL injection defence, role-based access, session control"),
    ]
    simple_table(doc, ["Capability", "Description"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 2. Technology Stack ────────────────────────────────────────────────────
    heading(doc, "2. Technology Stack")

    rows = [
        ("Odoo 16.0",         "Core ERP/CRM framework — provides ORM, HTTP, OWL frontend, job queue"),
        ("Python 3.10+",      "All custom business logic, models, controllers, services"),
        ("JavaScript / OWL",  "Odoo Web Library — frontend components (dashboards, charts, wizards)"),
        ("PostgreSQL 15",     "Primary relational database; materialized views for reporting"),
        ("PgBouncer",         "Connection pooler sitting between Odoo and Postgres (port 5433)"),
        ("Redis 7",           "Session storage, ORM read cache (ro_cache_redis), queue_job broker"),
        ("FastAPI",           "Fraud-detection microservice + RQ background workers/scheduler"),
        ("Apache SeaTunnel",  "ETL / data pipeline engine for multi-source data synchronisation"),
        ("Nginx",             "Reverse proxy, TLS termination, WebSocket upgrade routing"),
        ("Docker / Compose",  "Container orchestration for all services"),
        ("GitHub Actions",    "CI/CD — deploys to dev server on push to main/production"),
    ]
    simple_table(doc, ["Technology", "Role"], rows)
    doc.add_paragraph()

    heading(doc, "2.1  Third-Party Odoo Addons", level=2)
    rows = [
        ("queue_job (OCA)",          "Async job queue — used for heavy CSV imports, risk recalculation, ETL"),
        ("fpg_redis_session",        "Redis-backed session storage (replaces default file sessions)"),
        ("ro_cache_redis",           "Redis ORM read cache — dramatically reduces DB queries"),
        ("jazzy_backend_theme",      "Custom Odoo backend UI theme"),
        ("bi_sql_editor",            "In-app SQL query editor for analysts"),
        ("psql_query_execute",       "Secure parameterised PostgreSQL query execution from UI"),
        ("mail_composer_cc_bcc",     "Adds CC/BCC to Odoo mail composer"),
    ]
    simple_table(doc, ["Addon", "Purpose"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 3. Architecture ────────────────────────────────────────────────────────
    heading(doc, "3. System Architecture")

    heading(doc, "3.1  Infrastructure Overview", level=2)
    para(doc, (
        "All services run as Docker containers managed by a single docker-compose.yml. "
        "Nginx acts as the single entry point, terminates TLS, and routes traffic to the "
        "appropriate upstream service."
    ), space_after=6)
    arch_diagram_infra(doc)

    heading(doc, "3.2  Module Dependency Graph", level=2)
    para(doc, (
        "compliance_management is the central hub. All other custom modules depend on it. "
        "The correct installation order must be followed — see Section 6."
    ), space_after=6)
    arch_diagram_modules(doc)

    heading(doc, "3.3  Request Flow", level=2)
    arch_diagram_request_flow(doc)

    heading(doc, "3.4  Port Map", level=2)
    rows = [
        ("80",   "Nginx HTTP",              "Redirects to HTTPS"),
        ("443",  "Nginx HTTPS",             "Main entry point for all traffic"),
        ("8069", "Odoo HTTP",               "Internal — proxied by Nginx"),
        ("8072", "Odoo Longpolling / WS",   "Internal — proxied for /websocket, /longpolling/"),
        ("8073", "Custom WS Server",        "compliance_management CSV-import WebSocket"),
        ("8001", "FastAPI",                 "Internal — proxied at /api/"),
        ("5433", "PgBouncer",               "Odoo connects here, not directly to Postgres"),
        ("5432", "PostgreSQL",              "Internal only"),
        ("6379", "Redis",                   "Internal only"),
        ("5801", "SeaTunnel Hazelcast",     "Cluster port"),
        ("9080", "SeaTunnel REST API v2",   "ETL job submission"),
    ]
    simple_table(doc, ["Port", "Service", "Notes"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 4. Repository Structure ────────────────────────────────────────────────
    heading(doc, "4. Repository Structure")
    para(doc, "The repository root (~/matrics-app on the server) contains:", space_after=4)

    structure = [
        ("matrics_addons/",   "Custom iComply modules (primary development target)"),
        ("addons/",           "Third-party / OCA community addons"),
        ("custom/",           "Environment-specific overrides and FastAPI custom code"),
        ("fastapi/",          "FastAPI fraud-detection service source code"),
        ("seatunnel/",        "SeaTunnel job configs and YAML definitions"),
        ("etc/",              "odoo.conf, requirements.txt, odoo-server.log"),
        ("pgbouncer/",        "PgBouncer configuration and Dockerfile"),
        ("postgresql/",       "PostgreSQL data volume (gitignored)"),
        ("nginx.conf",        "Nginx reverse proxy configuration"),
        ("docker-compose.yml","Full service orchestration"),
        ("upgrade_modules.sh","Module install/update helper script"),
        ("init_db.sh",        "Database initialisation script"),
        ("docs/",             "Technical documentation (this document)"),
    ]
    simple_table(doc, ["Path", "Purpose"], structure)
    doc.add_paragraph()

    heading(doc, "4.1  Custom Module Structure (matrics_addons/)", level=2)
    rows = [
        ("compliance_management", "Central hub — KYC, risk, transactions, dashboards, PEP/sanctions"),
        ("alert_management",      "Rule-based alert engine, alert history, alert groups"),
        ("case_management",       "Investigation case workflows, exception management"),
        ("regulatory_reports",    "NFIU STR/CTR generation and submission"),
        ("nfiu_reporting",        "NFIU-specific data models extending regulatory_reports"),
        ("internal_control",      "Internal audit and control framework"),
        ("transaction_screening", "Real-time transaction rule screening"),
        ("etl_manager",           "Multi-database ETL connector (Postgres, MySQL, Oracle, Snowflake)"),
        ("rule_book / rule_book_api", "Policy rule definitions and REST API"),
        ("access_control",        "Extended RBAC, session control"),
        ("utility",               "Shared helpers used across modules"),
    ]
    simple_table(doc, ["Module", "Responsibility"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 5. Prerequisites ───────────────────────────────────────────────────────
    heading(doc, "5. Prerequisites")

    heading(doc, "5.1  Local Development Machine", level=2)
    bullet(doc, "OS: Ubuntu 20.04+ / macOS 12+ / Windows 11 with WSL2")
    bullet(doc, "Docker Desktop 4.x (or Docker Engine + Docker Compose plugin)")
    bullet(doc, "Git 2.x")
    bullet(doc, "VS Code (recommended) with Python, Odoo Snippets, Docker extensions")
    bullet(doc, "8 GB RAM minimum, 16 GB recommended (Odoo + Postgres + SeaTunnel are memory-heavy)")
    doc.add_paragraph()

    heading(doc, "5.2  Server (Production / Dev)", level=2)
    bullet(doc, "Ubuntu 22.04 LTS")
    bullet(doc, "Docker Engine 24+ and Docker Compose plugin")
    bullet(doc, "Nginx installed on the host (not in a container)")
    bullet(doc, "SSL certificate — wildcard *.novajii.com at /etc/ssl/wildcard_novajii/")
    bullet(doc, "Minimum 4 vCPU, 16 GB RAM, 100 GB SSD")
    doc.add_paragraph()

    heading(doc, "5.3  Credentials & Secrets", level=2)
    para(doc, "The following values must be set in the .env file before starting:", space_after=4)
    rows = [
        ("ODOO_PORT",      "Host-side HTTP port for Odoo (e.g. 10016)"),
        ("ODOO_GEVENT_PORT","Host-side longpolling port (e.g. 20016)"),
        ("ODOO_USER",      "PostgreSQL superuser username"),
        ("ODOO_PASSWORD",  "PostgreSQL superuser password"),
        ("STORAGE_PATH",   "Absolute path on host for persistent volumes"),
        ("REDIS_PORT",     "Redis host-side port (default 6379)"),
        ("FASTAPI_PORT",   "FastAPI host-side port (default 8001)"),
    ]
    simple_table(doc, ["Variable", "Description"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 6. Setup Guide ─────────────────────────────────────────────────────────
    heading(doc, "6. Step-by-Step Setup Guide")

    heading(doc, "Step 1 — Clone the repository", level=2)
    code_block(doc, "git clone git@github.com:novaji-introserve/matrics-app.git ~/matrics-app")
    code_block(doc, "cd ~/matrics-app")
    doc.add_paragraph()

    heading(doc, "Step 2 — Create the .env file", level=2)
    para(doc, "Copy the template and fill in all values:", space_after=2)
    code_block(doc, "cp .env.example .env   # ask the team for the actual values")
    code_block(doc, "nano .env")
    para(doc, "Required variables: ODOO_PORT, ODOO_GEVENT_PORT, ODOO_USER, ODOO_PASSWORD, STORAGE_PATH, REDIS_PORT, FASTAPI_PORT", size=10, color="555555")
    doc.add_paragraph()

    heading(doc, "Step 3 — Create storage directories", level=2)
    code_block(doc, "mkdir -p $STORAGE_PATH/odoo16/lib $STORAGE_PATH/odoo16/data")
    code_block(doc, "mkdir -p $STORAGE_PATH/postgresql $STORAGE_PATH/redis $STORAGE_PATH/seatunnel/logs")
    doc.add_paragraph()

    heading(doc, "Step 4 — Pull and start all services", level=2)
    code_block(doc, "docker compose pull")
    code_block(doc, "docker compose up -d")
    para(doc, "Wait ~60 seconds for PostgreSQL and PgBouncer to be ready before proceeding.", size=10, color="CC6600")
    doc.add_paragraph()

    heading(doc, "Step 5 — Initialise the database", level=2)
    code_block(doc, "chmod +x init_db.sh")
    code_block(doc, "./init_db.sh")
    para(doc, "This creates the Odoo database and installs the base module.", size=10, color="555555")
    doc.add_paragraph()

    heading(doc, "Step 6 — Install custom modules (in dependency order)", level=2)
    para(doc, "Use the upgrade script with the --full flag for a clean first install:", space_after=2)
    code_block(doc, "chmod +x upgrade_modules.sh")
    code_block(doc, "./upgrade_modules.sh --full")
    para(doc, "This installs modules in the correct sequence:", size=10, color="555555")
    for m in ["1. compliance_management", "2. alert_management", "3. case_management",
              "4. regulatory_reports", "5. nfiu_reporting"]:
        bullet(doc, m, level=1)
    para(doc, "Each module runs in its own Odoo process so all DB tables are committed before the next module starts.", size=10, color="CC6600")
    doc.add_paragraph()

    heading(doc, "Step 7 — Configure Nginx", level=2)
    code_block(doc, "sudo cp nginx.conf /etc/nginx/sites-available/matrics")
    code_block(doc, "sudo ln -s /etc/nginx/sites-available/matrics /etc/nginx/sites-enabled/")
    code_block(doc, "sudo nginx -t && sudo systemctl reload nginx")
    para(doc, "Ensure SSL cert exists at /etc/ssl/wildcard_novajii/_novajii_com.crt", size=10, color="555555")
    doc.add_paragraph()

    heading(doc, "Step 8 — Verify the deployment", level=2)
    numbered(doc, "Open https://dev1.novajii.com — you should see the Odoo login page")
    numbered(doc, "Log in as Administrator and navigate to Compliance → Dashboard")
    numbered(doc, "Check docker logs: docker compose logs -f odoo16")
    numbered(doc, "Check Nginx: sudo tail -f /var/log/nginx/matrics_error.log")
    doc.add_paragraph()
    doc.add_page_break()

    # ── 7. Day-to-Day Development ──────────────────────────────────────────────
    heading(doc, "7. Day-to-Day Development Workflow")

    heading(doc, "7.1  Installing / Updating a Module", level=2)
    code_block(doc, "# Install a specific module")
    code_block(doc, "./upgrade_modules.sh alert_management")
    code_block(doc, "")
    code_block(doc, "# Update multiple modules at once")
    code_block(doc, "./upgrade_modules.sh compliance_management alert_management")
    code_block(doc, "")
    code_block(doc, "# Full sequential install (fresh environment)")
    code_block(doc, "./upgrade_modules.sh --full")
    para(doc, "The script auto-detects whether to use -i (install) or -u (update) based on the module's DB state.", size=10, color="555555")
    doc.add_paragraph()

    heading(doc, "7.2  Restarting Services", level=2)
    code_block(doc, "# Restart only Odoo (after Python model change)")
    code_block(doc, "docker compose restart odoo16")
    code_block(doc, "")
    code_block(doc, "# Restart everything")
    code_block(doc, "docker compose down && docker compose up -d")
    doc.add_paragraph()

    heading(doc, "7.3  Viewing Logs", level=2)
    code_block(doc, "# Live Odoo logs")
    code_block(doc, "docker compose logs -f odoo16")
    code_block(doc, "")
    code_block(doc, "# Odoo log file inside container")
    code_block(doc, "tail -f etc/odoo-server.log")
    code_block(doc, "")
    code_block(doc, "# FastAPI logs")
    code_block(doc, "docker compose logs -f fastapi")
    doc.add_paragraph()

    heading(doc, "7.4  JavaScript / Frontend Changes", level=2)
    para(doc, (
        "OWL components live in static/src/components/<feature>/. "
        "Changes to JS/XML/CSS do NOT require a module update — a browser hard-refresh "
        "(Ctrl+Shift+R) is sufficient when dev_mode is active in odoo.conf."
    ), space_after=8)

    heading(doc, "7.5  Adding a New Field", level=2)
    numbered(doc, "Add the field in models/<file>.py")
    numbered(doc, "Add it to the view XML in views/<file>.xml")
    numbered(doc, "Update security/ir.model.access.csv if access control applies")
    numbered(doc, "Bump the module version in __manifest__.py")
    numbered(doc, "Run: ./upgrade_modules.sh <module_name>")
    doc.add_paragraph()
    doc.add_page_break()

    # ── 8. Key Configuration Files ─────────────────────────────────────────────
    heading(doc, "8. Key Configuration Files")

    heading(doc, "8.1  etc/odoo.conf", level=2)
    rows = [
        ("addons_path",          "/mnt/extra-addons, /mnt/matrics-addons, /mnt/custom/addons"),
        ("server_wide_modules",  "base, web, fpg_redis_session, queue_job, ro_cache_redis, jazzy_backend_theme"),
        ("workers",              "4 (multiprocessing mode)"),
        ("db_host / db_port",    "pgbouncer / 5433 (NOT direct Postgres)"),
        ("redis_host",           "redis (Docker service name)"),
        ("proxy_mode",           "True (required when behind Nginx)"),
        ("logfile",              "/etc/odoo/odoo-server.log"),
        ("admin_passwd",         "Master password for DB management — change before production"),
    ]
    simple_table(doc, ["Setting", "Value / Notes"], rows)
    doc.add_paragraph()

    heading(doc, "8.2  queue_job channels", level=2)
    para(doc, "Configured in odoo.conf under [queue_job]:", space_after=2)
    bullet(doc, "root:4 — up to 4 concurrent background jobs overall")
    bullet(doc, "root.compliance:2 — dedicated slots for risk scoring (never blocked by ETL)")
    doc.add_paragraph()

    heading(doc, "8.3  Redis configuration", level=2)
    rows = [
        ("Sessions",    "redis_session = True, TTL 604800s (7 days) auth / 86400s anon"),
        ("ORM cache",   "redis_cache_url = redis://redis:6379/0, TTL 3600s"),
        ("Queue jobs",  "Used as RQ broker for FastAPI background workers"),
    ]
    simple_table(doc, ["Use", "Details"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 9. Security Architecture ───────────────────────────────────────────────
    heading(doc, "9. Security Architecture")

    heading(doc, "9.1  SQL Injection Defence (5 layers)", level=2)
    para(doc, "All raw SQL execution (bi_sql_editor, psql_query_execute, dynamic charts) goes through a 5-layer defence:", space_after=4)
    numbered(doc, "Input validation — length, character set, null checks")
    numbered(doc, "Pattern detection — 170+ blocked patterns via sqlparse (DROP, DELETE, EXECUTE, etc.)")
    numbered(doc, "Header validation — checks for comment injections, stacked queries")
    numbered(doc, "Model-level constraints — @api.constrains on query fields")
    numbered(doc, "Parameterised execution — always self.env.cr.execute(sql, params) with statement_timeout")
    para(doc, "See matrics_addons/SQL_INJECTION_PROTECTION_IMPLEMENTATION_GUIDE.md for full details.", size=10, color="555555", space_after=6)

    heading(doc, "9.2  Role-Based Access Control", level=2)
    rows = [
        ("group_compliance_compliance_officer",         "Full compliance access"),
        ("group_compliance_branch_compliance_officer",  "Branch-scoped compliance access"),
        ("group_compliance_chief_compliance_officer",   "All compliance + configuration"),
        ("group_compliance_relationship_manager",       "KYC/customer onboarding"),
        ("group_compliance_transaction_monitoring_team","Transaction monitoring and alerts"),
        ("group_compliance_compliance_risk_manager",    "Risk assessment and universe"),
        ("base.group_system",                           "System admin — full access"),
    ]
    simple_table(doc, ["Group", "Scope"], rows)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 10. Troubleshooting ────────────────────────────────────────────────────
    heading(doc, "10. Common Issues & Troubleshooting")

    rows = [
        ("Module install crashes with\n'relation X does not exist'",
         "Fresh install — table not created yet. The init() methods in compliance_management are guarded with table-existence checks. Ensure you install in the correct sequence (Section 6)."),
        ("ParseError: External ID not found",
         "A data XML file is loaded before the file that defines the referenced record. Check __manifest__.py data list order — the defining file must appear first."),
        ("Module upgrade does nothing\n(no update log lines)",
         "The module is not installed in the DB. The upgrade script auto-detects this and uses -i. Confirm with: docker compose exec db psql -U odoo -d icomply_dev -c \"SELECT name, state FROM ir_module_module WHERE name='<module>';\""),
        ("Odoo redirects to /odoo/web",
         "Nginx server_name doesn't match the incoming hostname. Verify nginx.conf has the correct server_name and reload Nginx."),
        ("Chart validation warnings on install",
         "Expected — charts referencing tables that don't exist yet are skipped with a WARNING. They will validate correctly after tables are created."),
        ("NotNullViolation on target_model_id",
         "A chart record in a parent module references a model from a child module. Move the chart record to the child module's data file (see case_management/data/demo/case_charts.xml as the pattern)."),
        ("Container won't start — port in use",
         "Another process holds the port. Check: sudo ss -tlnp | grep <port>"),
        ("Redis connection refused",
         "Redis container not running. Check: docker compose ps redis"),
    ]
    simple_table(doc, ["Symptom", "Solution"], rows, header_bg="7B241C")
    doc.add_paragraph()
    doc.add_page_break()

    # ── 11. CI/CD Pipeline ─────────────────────────────────────────────────────
    heading(doc, "11. CI / CD Pipeline")
    para(doc, (
        "GitHub Actions (.github/workflows/deploy_dev.yml) deploys to the remote dev server "
        "on push to the main or production branches via SSH git pull. "
        "No automated module update is run by CI — module upgrades must be triggered manually "
        "on the server after each deploy using upgrade_modules.sh."
    ), space_after=6)

    heading(doc, "Deployment checklist after a push:", level=2)
    numbered(doc, "Push to main branch (GitHub Actions runs git pull on server automatically)")
    numbered(doc, "SSH into the server: ssh ubuntu@dev1.novajii.com")
    numbered(doc, "cd ~/matrics-app")
    numbered(doc, "Run the upgrade for changed modules: ./upgrade_modules.sh <module>")
    numbered(doc, "Restart Odoo if Python files changed: docker compose restart odoo16")
    numbered(doc, "Verify: docker compose logs --tail=50 odoo16")
    doc.add_paragraph()
    doc.add_page_break()

    # ── 12. Contacts / References ──────────────────────────────────────────────
    heading(doc, "12. Contacts & References")

    heading(doc, "12.1  Internal Contacts", level=2)
    rows = [
        ("Engineering Lead",      "Jonathan Ogbimi", "jonathan.o@novajii.com"),
        ("Backend Developer",     "TBD",             "—"),
    ]
    simple_table(doc, ["Role", "Name", "Email"], rows)
    doc.add_paragraph()

    heading(doc, "12.2  Key References", level=2)
    bullet(doc, "Odoo 16.0 docs: https://www.odoo.com/documentation/16.0/")
    bullet(doc, "Odoo 16.0 source (local): /home/jonathan/projects/odoo/16/src/odoo-16.0/")
    bullet(doc, "OCA queue_job: https://github.com/OCA/queue")
    bullet(doc, "Apache SeaTunnel 2.3.12: https://seatunnel.apache.org/docs/2.3.12/")
    bullet(doc, "Docker Compose reference: https://docs.docker.com/compose/")
    bullet(doc, "PgBouncer docs: https://www.pgbouncer.org/config.html")
    doc.add_paragraph()

    divider(doc)
    para(doc, "© 2026 Novaji Introserve Limited — Confidential", size=9,
         color="888888", align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── save ───────────────────────────────────────────────────────────────────
    out = "/home/jonathan/projects/matrics/docs/iComply_MATRICS_KTD.docx"
    doc.save(out)
    print(f"Document saved to: {out}")


if __name__ == "__main__":
    build()
